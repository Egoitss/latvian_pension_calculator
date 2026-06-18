import io
import os
import sys

# Allow importing from the project root
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import yaml
from docx import Document

import download_counter
from export_docx import build_summary_docx
from i18n import make_t
import app as app_module


SAMPLE = {
    "inputs": {
        "age": 30, "retirementAge": 65,
        "grossMonthly": 1750, "scenario": "moderate",
    },
    "pillars": {
        "p1": {"monthly": 412, "capital": 120000},
        "p2": {"monthly": 689, "capital": 210000},
        "p3": {"monthly": 210, "capital": 60000},
    },
    "totals": {
        "monthly": 1311, "realMonthly": 980,
        "capital": 390000, "realCapital": 250000,
    },
}


def test_lv_yaml_parses_without_duplicate_keys():
    # The catalog must load cleanly after the new export strings.
    path = os.path.join(
        os.path.dirname(__file__), "..", "translations", "lv.yaml")
    data = yaml.safe_load(open(path, encoding="utf-8"))
    assert isinstance(data, dict)
    assert data["Pension summary"] == "Pensijas kopsavilkums"


def test_build_docx_returns_openable_document():
    blob = build_summary_docx(SAMPLE, make_t("en"), "2026-06-17")
    assert isinstance(blob, bytes) and len(blob) > 1000
    doc = Document(io.BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Pension summary" in text
    assert "1 311" in text          # space-thousands EUR formatting


def test_build_docx_localizes_to_latvian():
    blob = build_summary_docx(SAMPLE, make_t("lv"), "2026-06-17")
    text = "\n".join(
        p.text for p in Document(io.BytesIO(blob)).paragraphs)
    assert "Pensijas kopsavilkums" in text


def test_counter_increments_and_persists(tmp_path, monkeypatch):
    monkeypatch.setenv(
        "DOWNLOAD_COUNT_FILE", str(tmp_path / "count.json"))
    assert download_counter.read_count() == 0
    assert download_counter.increment() == 1
    assert download_counter.increment() == 2
    assert download_counter.read_count() == 2


def _client():
    app_module.app.config["TESTING"] = True
    return app_module.app.test_client()


def test_export_route_returns_docx_and_counts(tmp_path, monkeypatch):
    monkeypatch.setenv(
        "DOWNLOAD_COUNT_FILE", str(tmp_path / "count.json"))
    resp = _client().post("/export/docx", json=SAMPLE)
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["Content-Type"]
    assert "attachment" in resp.headers["Content-Disposition"]
    assert resp.headers["X-Download-Count"] == "1"
    assert len(resp.data) > 1000


def test_download_count_endpoint(tmp_path, monkeypatch):
    monkeypatch.setenv(
        "DOWNLOAD_COUNT_FILE", str(tmp_path / "count.json"))
    client = _client()
    client.post("/export/docx", json=SAMPLE)
    client.post("/lv/export/docx", json=SAMPLE)
    data = client.get("/api/download-count").get_json()
    assert data["count"] == 2


def test_export_route_survives_empty_payload(tmp_path, monkeypatch):
    monkeypatch.setenv(
        "DOWNLOAD_COUNT_FILE", str(tmp_path / "count.json"))
    resp = _client().post("/export/docx", json={})
    assert resp.status_code == 200
    assert len(resp.data) > 500
