# export_docx.py — render the calculator state into a one-page,
# branded Word (.docx) summary. Pure formatting; all numbers come
# from the client payload, all labels via the t() translator.
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

# Brand palette (mirrors the app's slate/emerald Tailwind tones)
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
SLATE_400 = RGBColor(0x94, 0xA3, 0xB8)
EMERALD = RGBColor(0x05, 0x96, 0x69)


def _fmt_eur(value) -> str:
    # "€ 1 311" — space thousands, matching the lv-LV UI formatting.
    try:
        n = round(float(value))
    except (TypeError, ValueError):
        n = 0
    return "€ " + f"{n:,.0f}".replace(",", " ")


def _line(doc, text, size=10, color=SLATE_500, bold=False, after=2):
    # Add a single-run paragraph with tight spacing.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p


def _set_margins(doc):
    # Narrow margins keep the whole summary on a single page.
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


def _add_header(doc, t, date_str):
    # Title row + date.
    _line(doc, t("Pension summary"), size=20, color=SLATE_900,
          bold=True, after=0)
    _line(doc, date_str, size=9, color=SLATE_400, after=10)


def _add_hero(doc, t, totals):
    # Headline monthly figure + today's-money line.
    _line(doc, t("Your retirement income"), size=9, color=SLATE_400,
          after=0)
    # Lead with real (today's money); nominal is the secondary line.
    _line(doc,
          _fmt_eur(totals.get("realMonthly", 0)) + " " + t("/month"),
          size=24, color=EMERALD, bold=True, after=0)
    _line(doc,
          t("nominal — future euros, before inflation") + ": "
          + _fmt_eur(totals.get("monthly", 0)) + " " + t("/month"),
          size=11, color=SLATE_500, after=10)


def _add_table(doc, t, pillars, totals):
    # Per-pillar monthly + capital, then a totals row.
    rows = [
        (t("State pension"), "p1"),
        (t("Investment pension"), "p2"),
        (t("Voluntary pension"), "p3"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _cell(hdr[0], t("Pillar"), bold=True)
    _cell(hdr[1], t("Monthly"), bold=True, align="right")
    _cell(hdr[2], t("Capital at retirement"), bold=True, align="right")
    for label, key in rows:
        d = pillars.get(key, {})
        cells = table.add_row().cells
        _cell(cells[0], label)
        _cell(cells[1], _fmt_eur(d.get("monthly", 0)), align="right")
        _cell(cells[2], _fmt_eur(d.get("capital", 0)), align="right")
    tot = table.add_row().cells
    _cell(tot[0], t("Total"), bold=True)
    _cell(tot[1], _fmt_eur(totals.get("monthly", 0)), bold=True,
          align="right")
    _cell(tot[2], _fmt_eur(totals.get("capital", 0)), bold=True,
          align="right")


def _cell(cell, text, bold=False, align="left"):
    # Write one table cell with compact, branded text.
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if align == "right"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    run.font.color.rgb = SLATE_900 if bold else SLATE_500


def _add_inputs(doc, t, inputs):
    # One-line recap of the assumptions behind the numbers.
    parts = [
        f"{t('Age')}: {inputs.get('age', '—')}",
        f"{t('Retirement age')}: {inputs.get('retirementAge', '—')}",
        f"{t('Gross monthly')}: {_fmt_eur(inputs.get('grossMonthly', 0))}",
        f"{t('Scenario')}: {inputs.get('scenario', '—')}",
    ]
    _line(doc, " · ".join(parts), size=9, color=SLATE_400, after=8)


def build_summary_docx(data, t, date_str="") -> bytes:
    # Assemble the document and return the .docx bytes.
    doc = Document()
    _set_margins(doc)
    _add_header(doc, t, date_str)
    _add_hero(doc, t, data.get("totals", {}))
    _add_table(doc, t, data.get("pillars", {}), data.get("totals", {}))
    _add_inputs(doc, t, data.get("inputs", {}))
    _line(doc,
          t("Simulation model — not financial advice."),
          size=8, color=SLATE_400, after=0)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
